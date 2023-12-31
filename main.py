from PyQt6 import QtWidgets
from PyQt6.QtWidgets import *
import sys
import sqlite3
import time
import os 
import subprocess
import docx
import random
OPERATING_SYSTEM = "mac"
cursor = None
connection = None
tests_directory = "tests"
parent_dir = os.getcwd()
path = os.path.join(parent_dir, tests_directory) 


def check_db_connection(connection):
    try:
        connection.cursor()
        return True
    except Exception as ex:
        return False
    
def create_new_tables():
    global cursor
    global connection
    cursor.execute("CREATE TABLE test(id, nazwa, kategoria, subkategoria, liczba_pytan);")
    cursor.execute("CREATE TABLE test_pytanie(id_testu, id_pytania);")
    cursor.execute("CREATE TABLE pytanie(id, kategoria, subkategoria, pytanie, odpowiedzA, odpowiedzB, odpowiedzC, odpowiedzD, is_correct);")
    cursor.execute("CREATE TABLE nadkategoria(id, nazwa);")
    cursor.execute("CREATE TABLE subkategoria(id, nazwa, nadkategoria_nazwa);")
    connection.commit()
    cursor.close()

def get_latest_id(table):
    global connection
    global cursor
    
    cursor.execute("SELECT id FROM " + table + " ORDER BY id DESC LIMIT 1;")
    connection.commit()
    result = cursor.fetchone()
    if result == None:
        return 0
    number = int(result[0])
    return number
    

def get_categories_names():
    global connection
    global cursor

    connection = sqlite3.connect("test-gen-db.db")
    categories = []
    if check_db_connection(connection):
        cursor = connection.cursor()
        cursor.execute("SELECT nazwa FROM nadkategoria;")
        connection.commit()
        result = cursor.fetchall()
        for item in result:
            categories.append(item[0])
    
    cursor.close()
    connection.close()
    return categories


def get_subcategories_names(category_name):
    global connection
    global cursor

    connection = sqlite3.connect("test-gen-db.db")
    subcategories = []
    if check_db_connection(connection):
        cursor = connection.cursor()
        cursor.execute("SELECT nazwa FROM subkategoria WHERE nadkategoria_nazwa = '" + category_name + "';")
        connection.commit()
        result = cursor.fetchall()
        for item in result:
            subcategories.append(item[0])

    cursor.close()
    connection.close()
    return subcategories


def get_all_tests():
    global connection
    global cursor

    connection = sqlite3.connect("test-gen-db.db")
    tests = []
    if check_db_connection(connection):
        cursor = connection.cursor()
        cursor.execute("SELECT id, nazwa FROM test;")
        connection.commit()
        result = cursor.fetchall()
        for item in result:
            item_2 = str(item[0]) + "_" + str(item[1])
            tests.append(item_2.replace(" ", "_"))

    cursor.close()
    connection.close()
    return tests


def get_all_questions_from_subcategory(category, subcategory):
    global connection
    global cursor

    questions = []
    cursor.execute("SELECT * FROM pytanie WHERE kategoria = '" + category 
                       + "' AND subkategoria = '" + subcategory + "';")
    result = cursor.fetchall()
    for item in result:
        questions.append(item)
        
    return questions


def get_all_selected_questions(selected_questions):
    global connection
    global cursor

    questions = []
    for question_id in selected_questions:
        cursor.execute("SELECT * FROM pytanie WHERE id = " + str(question_id) + ";")
        result = cursor.fetchone()
        questions.append(result)

    return questions


class MainAppWindow(QMainWindow):
    def __init__(self):
        super(MainAppWindow, self).__init__()
        self.setGeometry(100, 100, 1000, 600)
        self.setWindowTitle("Document creator")
        self.main_layout = QGridLayout(self)
        self.setLayout(self.main_layout)
        self.initUI()

    def initUI(self):

        self.tab_widget = QtWidgets.QTabWidget(self)
        self.tab_widget.setGeometry(0, 0, 1000, 600)
        
        self.widget_add_to_db = QtWidgets.QWidget(self)
        self.tab1UI()

        self.widget_add_category = QtWidgets.QWidget(self)
        self.tab2UI()

        self.widget_show_test_menu = QtWidgets.QWidget(self)
        self.tab3UI()

        self.widget_generate_test_menu = QtWidgets.QWidget(self)
        self.tab4UI()

        self.tab_widget.addTab(self.widget_add_to_db, "Dodaj do bazy")
        self.tab_widget.addTab(self.widget_add_category, "Dodaj kategorię")
        self.tab_widget.addTab(self.widget_show_test_menu, "Wyświetl test")
        self.tab_widget.addTab(self.widget_generate_test_menu, "Wygeneruj test")

        self.main_layout.addWidget(self.tab_widget, 0, 0, 2, 1)
        self.show()

    def tab1UI(self):

        self.layout = QGridLayout()
        self.widget_add_to_db.setLayout(self.layout)

        self.combo_category = QComboBox(self)
        self.combo_category.setFixedWidth(500)
        self.layout.addWidget(QLabel("Wpisz kategorię pytania: "), 0, 0)
        self.layout.addWidget(self.combo_category, 0, 1)
        self.categories = get_categories_names()
        for category in self.categories:
            self.combo_category.addItem(category)

        self.combo_subcategory = QComboBox(self)
        self.combo_subcategory.setFixedWidth(500)
        self.layout.addWidget(QLabel("Wpisz podkategorię pytania: "), 1, 0)
        self.layout.addWidget(self.combo_subcategory, 1, 1)
        self.subcategories = get_subcategories_names(str(self.combo_category.currentText()))
        for subcategory in self.subcategories:
            self.combo_subcategory.addItem(subcategory)
        
        self.combo_category.currentTextChanged.connect(self.update_subcategory_combobox)

        self.question = QtWidgets.QPlainTextEdit(self)
        self.question.setFixedWidth(500)
        self.question.setFixedHeight(100)
        self.layout.addWidget(QLabel("Wpisz pytanie: "), 2, 0)
        self.layout.addWidget(self.question, 2, 1)
        
        self.answerA = QtWidgets.QPlainTextEdit(self)
        self.answerA.setFixedWidth(500)
        self.layout.addWidget(QLabel("Wpisz odpowiedź A: "), 3, 0)
        self.layout.addWidget(self.answerA, 3, 1)

        self.is_answerA_correct_cbx = QCheckBox("")
        self.is_answerA_correct_cbx.setChecked(False)
        self.layout.addWidget(self.is_answerA_correct_cbx, 3, 2)

        self.answerB = QtWidgets.QPlainTextEdit(self)
        self.answerB.setFixedWidth(500)
        self.layout.addWidget(QLabel("Wpisz odpowiedź B: "), 4, 0)
        self.layout.addWidget(self.answerB, 4, 1)

        self.is_answerB_correct_cbx = QCheckBox("")
        self.is_answerB_correct_cbx.setChecked(False)
        self.layout.addWidget(self.is_answerB_correct_cbx, 4, 2)

        self.answerC = QtWidgets.QPlainTextEdit(self)
        self.answerC.setFixedWidth(500)
        self.layout.addWidget(QLabel("Wpisz odpowiedź C: "), 5, 0)
        self.layout.addWidget(self.answerC, 5, 1)

        self.is_answerC_correct_cbx = QCheckBox("")
        self.is_answerC_correct_cbx.setChecked(False)
        self.layout.addWidget(self.is_answerC_correct_cbx, 5, 2)

        self.answerD = QtWidgets.QPlainTextEdit(self)
        self.answerD.setFixedWidth(500)
        self.layout.addWidget(QLabel("Wpisz odpowiedź D: "), 6, 0)
        self.layout.addWidget(self.answerD, 6, 1)

        self.is_answerD_correct_cbx = QCheckBox("")
        self.is_answerD_correct_cbx.setChecked(False)
        self.layout.addWidget(self.is_answerD_correct_cbx, 6, 2)

        self.add_question_btn = QPushButton()
        self.add_question_btn.setObjectName("add question")
        self.add_question_btn.setText("Dodaj pytanie")
        self.layout.addWidget(QLabel("Dodaj pytanie do bazy pytań: "), 7, 0)
        self.layout.addWidget(self.add_question_btn, 7, 1)

        self.add_question_btn.clicked.connect(self.add_new_question)

    def get_correct_answer(self):
        if (self.is_answerA_correct_cbx.isChecked()):
            return "odpowiedzA"
        if (self.is_answerB_correct_cbx.isChecked()):
            return "odpowiedzB"
        if (self.is_answerC_correct_cbx.isChecked()):
            return "odpowiedzC"
        if (self.is_answerD_correct_cbx.isChecked()):
            return "odpowiedzD"
        return False

    def tab2UI(self):
        self.layout = QFormLayout()
        self.widget_add_category.setLayout(self.layout)

        self.main_category = QtWidgets.QLineEdit(self)
        self.main_category.setFixedWidth(500)
        self.layout.addRow("Dodaj główną kategorię: ", self.main_category)

        self.add_main_category_btn = QPushButton()
        self.add_main_category_btn.setObjectName("add main category")
        self.add_main_category_btn.setText("Dodaj kategorię")
        self.layout.addRow("Dodaj kategorię do bazy kategorii: ", self.add_main_category_btn)

        self.sub_category = QtWidgets.QLineEdit(self)
        self.sub_category.setFixedWidth(500)
        self.layout.addRow("Dodaj podkategorię: ", self.sub_category)

        self.combo_supercategory = QComboBox(self)
        self.combo_supercategory.setFixedWidth(500)
        self.layout.addRow("Wybierz nadkategorię: ", self.combo_supercategory)
        self.supercategories = get_categories_names()
        for category in self.supercategories:
            self.combo_supercategory.addItem(category)

        self.add_sub_category_btn = QPushButton()
        self.add_sub_category_btn.setObjectName("add sub category")
        self.add_sub_category_btn.setText("Dodaj podkategorię")
        self.layout.addRow("Dodaj podkategorię do bazy podkategorii: ", self.add_sub_category_btn)

        self.add_main_category_btn.clicked.connect(self.add_main_category)
        self.add_main_category_btn.clicked.connect(self.update_category_combobox)
        self.add_main_category_btn.clicked.connect(self.update_category_combobox2)
        self.add_main_category_btn.clicked.connect(self.update_category_combobox3)

        self.add_sub_category_btn.clicked.connect(self.add_sub_category)
        self.add_sub_category_btn.clicked.connect(self.update_subcategory_combobox)
        self.add_sub_category_btn.clicked.connect(self.update_subcategory_combobox2)


    def tab3UI(self):
        self.layout = QFormLayout()
        self.widget_show_test_menu.setLayout(self.layout)
    
        self.combo_test_name = QComboBox(self)
        self.combo_test_name.setFixedWidth(500)
        self.layout.addRow("Nazwa testu: ", self.combo_test_name)
        self.tests = get_all_tests()
        for test in self.tests:
            self.combo_test_name.addItem(test)

        self.show_test_btn = QPushButton()
        self.show_test_btn.setObjectName("show test")
        self.show_test_btn.setText("Pokaż test")
        self.layout.addRow("Wczytaj test z bazy testów: ", self.show_test_btn)

        self.show_test_btn.clicked.connect(self.read_test)

        
    def tab4UI(self):
        self.layout = QFormLayout()
        self.widget_generate_test_menu.setLayout(self.layout)
        self.test_name = QtWidgets.QLineEdit(self)
        self.test_name.setFixedWidth(500)
        self.layout.addRow("Nazwa testu: ", self.test_name)

        self.combo_category_of_questions = QComboBox(self)
        self.combo_category_of_questions.setFixedWidth(500)
        self.layout.addRow("Kategoria pytań: ", self.combo_category_of_questions)
        self.supercategories = get_categories_names()
        for category in self.supercategories:
            self.combo_category_of_questions.addItem(category)

        self.combo_subcategory_of_questions = QComboBox(self)
        self.combo_subcategory_of_questions.setFixedWidth(500)
        self.layout.addRow(QLabel("Podkategoria pytania: "), self.combo_subcategory_of_questions)
        self.subcategories = get_subcategories_names(str(self.combo_category.currentText()))
        for subcategory in self.subcategories:
            self.combo_subcategory_of_questions.addItem(subcategory)
        
        self.combo_category_of_questions.currentTextChanged.connect(self.update_subcategory_combobox2)

        self.number_of_questions = QtWidgets.QLineEdit(self)
        self.number_of_questions.setFixedWidth(500)
        self.layout.addRow("Liczba pytań: ", self.number_of_questions)

        self.generate_test_btn = QPushButton()
        self.generate_test_btn.setObjectName("generate test")
        self.generate_test_btn.setText("Wygeneruj test")
        self.layout.addRow("Dodaj test do bazy testów: ", self.generate_test_btn)

        self.generate_test_btn.clicked.connect(self.generate_test)
        self.generate_test_btn.clicked.connect(self.update_tests_combobox)
        

    def check_is_empty_tab1UI(self):
        if (str(self.combo_category.currentText()) == "" or str(self.combo_subcategory.currentText()) == ""
            or self.question.toPlainText() == "" or self.answerA.toPlainText() == "" or self.answerB.toPlainText() == ""
            or self.answerC.toPlainText() == "" or self.answerD.toPlainText() == ""):
            return True
        return False

    def check_is_empty_tab3UI(self):
        pass

    def check_is_empty_tab4UI(self):
        if (self.test_name.text() == "" or str(self.combo_category_of_questions.currentText()) == "" or
            str(self.combo_subcategory_of_questions.currentText()) == "" or
            self.number_of_questions.text() == ""):
            return True
        return False

    def info_not_all_fields(self):
        QMessageBox.information(self, 'Nie wypełniono wszystkich pól!', 'Wypełnij wszystkie pola.')

    def info_correct_add_question(self):
        QMessageBox.information(self, 'Poprawnie dodano pytanie do bazy danych!', 'Mozesz dodać kolejne pytanie.')

    def info_incorrect_add_question(self):
        QMessageBox.warning(self, 'Nie dodano pytania do bazy danych!', 'Błąd dodawania pytania do bazy.')

    def info_correct_add_category(self):
        QMessageBox.information(self, 'Poprawnie dodano kategorię do bazy danych!', 'Mozesz dodać kolejną kategorię.')

    def info_incorrect_add_category(self):
        QMessageBox.warning(self, 'Nie dodano kategorii do bazy danych!', 'Błąd dodawania kategorii do bazy.')

    def info_correct_add_test(self):
        QMessageBox.information(self, 'Poprawnie dodano test do bazy danych!', 'Mozesz dodać kolejny test.')

    def info_incorrect_add_test(self):
        QMessageBox.warning(self, 'Nie dodano testu do bazy danych!', 'Błąd dodawania testu do bazy.')

    def info_incorrect_number_of_questions(self):
        QMessageBox.warning(self, 'Za mało pytań w bazie pytań w podanej podkategorii!', 'Dodaj więcej pytań.')

    def add_new_question(self):
        global cursor
        global connection

        is_empty = self.check_is_empty_tab1UI()
        if is_empty:
            self.info_not_all_fields()

        else:
            connection = sqlite3.connect("test-gen-db.db")
            if check_db_connection(connection):
                cursor = connection.cursor()
                print("Dodawanie pytania do bazy danych")
                id = get_latest_id("pytanie")     
                if id != -1:
                    current_id = id + 1
                    insert_statement = """INSERT INTO pytanie 
                                        (id, kategoria, subkategoria, pytanie, odpowiedzA, odpowiedzB, odpowiedzC, odpowiedzD, is_correct) 
                                        VALUES ({0}, "{1}", "{2}", "{3}", "{4}", "{5}", "{6}", "{7}", "{8}");""".format(
                                            current_id,
                                            self.combo_category.currentText(),
                                            self.combo_subcategory.currentText(),
                                            self.question.toPlainText(),
                                            self.answerA.toPlainText(),
                                            self.answerB.toPlainText(),
                                            self.answerC.toPlainText(),
                                            self.answerD.toPlainText(),
                                            self.get_correct_answer()
                                        )
                    print(insert_statement)
                    cursor.execute(insert_statement)
                    time.sleep(5)
                    connection.commit()

                else:
                    print("Błąd wczytania id z bazy danych pytanie")
                
                inserted_id = get_latest_id("pytanie")
                if inserted_id == id + 1:
                    print("Pytanie zostało dodane prawidłowo")
                    self.info_correct_add_question()
                    self.question.setPlainText("")
                    self.answerA.setPlainText("")
                    self.answerB.setPlainText("")
                    self.answerC.setPlainText("")
                    self.answerD.setPlainText("")

                else:
                    print("Pytanie nie zostało dodane prawidłowo")
                    self.info_incorrect_add_question()

            else:
                print("Błąd dodawania do bazy danych: brak połączenia")
                self.info_incorrect_add_question()
            
            

    def update_category_combobox(self):
        self.combo_category.clear()
        self.categories = get_categories_names()
        for category in self.categories:
            self.combo_category.addItem(category)


    def update_subcategory_combobox(self):
        self.combo_subcategory.clear()
        self.subcategories = get_subcategories_names(str(self.combo_category.currentText()))
        for subcategory in self.subcategories:
            self.combo_subcategory.addItem(subcategory)


    def update_category_combobox2(self):
        self.combo_category_of_questions.clear()
        self.categories = get_categories_names()
        for category in self.categories:
            self.combo_category_of_questions.addItem(category)
    

    def update_subcategory_combobox2(self):
        self.combo_subcategory_of_questions.clear()
        self.subcategories = get_subcategories_names(str(self.combo_category_of_questions.currentText()))
        for subcategory in self.subcategories:
            self.combo_subcategory_of_questions.addItem(subcategory)


    def update_category_combobox3(self):
        self.combo_supercategory.clear()
        self.categories = get_categories_names()
        for category in self.categories:
            self.combo_supercategory.addItem(category)


    def update_tests_combobox(self):
        self.combo_test_name.clear()
        self.tests = get_all_tests()
        for test in self.tests:
            self.combo_test_name.addItem(test)


    def add_main_category(self):
        global cursor
        global connection

        if (self.main_category.text() == ""):
            self.info_not_all_fields()
        else:
            connection = sqlite3.connect("test-gen-db.db")
            if check_db_connection(connection):
                cursor = connection.cursor()
                id = get_latest_id("nadkategoria")
                if id != -1:
                    current_id = id + 1
                    insert_statement = """INSERT INTO nadkategoria
                                        (id, nazwa) 
                                        VALUES ({0}, "{1}");""".format(
                                            current_id,
                                            self.main_category.text()
                                        )
                    print(insert_statement)
                    cursor.execute(insert_statement)
                    time.sleep(5)
                    connection.commit()

                else:
                    print("Błąd wczytania id z bazy danych pytanie")
                
                inserted_id = get_latest_id("nadkategoria")
                if inserted_id == id + 1:
                    print("Kategoria została dodana prawidłowo")
                    self.info_correct_add_category()
                    self.main_category.setText("")
                    self.update_subcategory_combobox()
                    self.update_subcategory_combobox2()
                else:
                    print("Kategoria nie została dodana prawidłowo")
                    self.info_incorrect_add_category()

            else:
                print("Błąd dodawania do bazy danych: brak połączenia")
                self.info_incorrect_add_category()



    def add_sub_category(self):
        global cursor
        global connection

        if (self.sub_category.text() == ""):
            self.info_not_all_fields()
        else:
            connection = sqlite3.connect("test-gen-db.db")
            if check_db_connection(connection):
                cursor = connection.cursor()
                id = get_latest_id("subkategoria")
                if id != -1:
                    current_id = id + 1
                    insert_statement = """INSERT INTO subkategoria
                                        (id, nazwa, nadkategoria_nazwa) 
                                        VALUES ({0}, "{1}", "{2}");""".format(
                                            current_id,
                                            self.sub_category.text(),
                                            self.combo_supercategory.currentText()
                                        )
                    print(insert_statement)
                    cursor.execute(insert_statement)
                    time.sleep(5)
                    connection.commit()

                else:
                    print("Błąd wczytania id z bazy danych subkategoria")
                
                inserted_id = get_latest_id("subkategoria")
                if inserted_id == id + 1:
                    print("Subkategoria została dodana prawidłowo")
                    self.info_correct_add_category()
                    self.sub_category.setText("")
                else:
                    print("Subkategoria nie została dodana prawidłowo")
                    self.info_incorrect_add_category()

            else:
                print("Błąd dodawania do bazy danych: brak połączenia")
                self.info_incorrect_add_category()


    def read_test(self):
        current_path =  path + "/" + str(self.combo_test_name.currentText()) + ".docx"
        if OPERATING_SYSTEM == "mac":
            subprocess.call(('open', current_path))

        if OPERATING_SYSTEM == "windows":
            os.startfile(current_path)

    def generate_test(self):
        global cursor
        global connection

        is_empty = self.check_is_empty_tab4UI()
        if is_empty:
            self.info_not_all_fields()
        else:
            connection = sqlite3.connect("test-gen-db.db")
            if check_db_connection(connection):
                cursor = connection.cursor()
                print("Dodawanie testu do bazy testów")
                id = get_latest_id("test")
                current_id = -1
                if id != -1:
                    current_id = id + 1
                    is_generated = self.generate_questions(current_id)
                    if is_generated:
                        insert_statement = """INSERT INTO test
                        (id, nazwa, kategoria, subkategoria, liczba_pytan)
                        VALUES ({0}, "{1}", "{2}", "{3}", {4});""".format(      
                            current_id,
                            self.test_name.text(),
                            str(self.combo_category_of_questions.currentText()),
                            str(self.combo_subcategory_of_questions.currentText()),
                            self.number_of_questions.text()
                        )
                        print(insert_statement)
                        cursor.execute(insert_statement)
                        time.sleep(5)
                        connection.commit()
                else:
                    print("Błąd wczytania id z bazy danych test")
            
                inserted_id = get_latest_id("test")
                if inserted_id == id + 1:
                    
                    print("Test został dodany prawidłowo")
                    self.info_correct_add_test()
                    self.test_name.setText("")
                    self.number_of_questions.setText("")
                    
                else:
                    print("Test nie został dodany prawidłowo")
                    self.info_incorrect_add_test()

    
            else:
                print("Błąd dodawania do bazy danych: brak połączenia")
                self.info_incorrect_add_test()


    def generate_questions(self, test_id):
        category = str(self.combo_category_of_questions.currentText())
        subcategory = str(self.combo_subcategory_of_questions.currentText())
        number_of_questions = self.number_of_questions.text()

        questions = get_all_questions_from_subcategory(category, subcategory)
        questions_ids = []
        for question in questions:
            questions_ids.append(question[0])

        print(questions)
        print(questions_ids)

        if len(questions) < int(number_of_questions):
            print("Za mało pytań w podkategorii by ułożyć test")
            self.info_incorrect_number_of_questions()
            return False
        else:
            print("Wystarczająca liczba pytań by ułożyć test")
            selected_questions = []
            i = 0
            while i < len(questions_ids):
                selected = random.choice(questions_ids)
                while selected in selected_questions:
                    selected = random.choice(questions_ids)
                selected_questions.append(selected)
                i += 1

            print("SELECTED:" + str(selected_questions))
            for question_id in selected_questions:
                insert_statement = """INSERT INTO test_pytanie
                (id_testu, id_pytania)
                VALUES ({0}, {1});""".format(      
                    test_id,
                    question_id
                )
                print(insert_statement)
                cursor.execute(insert_statement)
                time.sleep(5)
                connection.commit()

            self.write_questions_to_docx(test_id, selected_questions)
            return True
            

    def write_questions_to_docx(self, test_id, selected_questions):

        doc = docx.Document()
        test_temp = self.test_name.text()
        test_name_replaced = test_temp.replace(" ", "_")
        docx_path = "tests/" + str(test_id) + "_" + test_name_replaced + ".docx"

        doc_answers = docx.Document()
        docx_path_answers = "tests/" + str(test_id) + "_" + test_name_replaced + "_answers" + ".docx"
        
        questions = get_all_selected_questions(selected_questions)
        for index, question in enumerate(questions):
            question_text = question[3]

            paragraph = doc.add_paragraph()
            run = paragraph.add_run(str(index+1) + ". " + question_text)
            run.font.name = 'Arial'
            run.font.size = docx.shared.Pt(12)

            paragraph_answ = doc_answers.add_paragraph()
            run_answ = paragraph_answ.add_run(str(index+1) + ". " + question_text)
            run_answ.font.name = 'Arial'
            run_answ.font.size = docx.shared.Pt(12)
    
            correct_answer_text = -1
            rest_answers = []
            correct_answer = question[8]
            t = [0, 1, 2, 3]
            choice = random.choice(t)
            if correct_answer == 'odpowiedzA':
                correct_answer_text = question[4]
                rest_answers = [question[5], question[6], question[7]]

            elif correct_answer == 'odpowiedzB':
                correct_answer_text = question[5]
                rest_answers = [question[4], question[6], question[7]]

            elif correct_answer == 'odpowiedzC':
                correct_answer_text = question[6]
                rest_answers = [question[4], question[5], question[7]]

            elif correct_answer == 'odpowiedzD':
                correct_answer_text = question[7]
                rest_answers = [question[4], question[5], question[6]]
    
            random.shuffle(rest_answers)
            rest_answers.insert(choice, correct_answer_text)

            paragraph = doc.add_paragraph()
            run = paragraph.add_run("A: " + rest_answers[0] + " B: " + rest_answers[1] + " C: " + 
                                    rest_answers[2] + " D: " + rest_answers[3])
            run.font.name = 'Arial'
            run.font.size = docx.shared.Pt(12)

            paragraph_answ = doc_answers.add_paragraph()
            if choice == 0:
                run_answ = paragraph_answ.add_run("Poprawna odp. : A (" + rest_answers[0] + ")")

            elif choice == 1:
                run_answ = paragraph_answ.add_run("Poprawna odp. : B (" + rest_answers[1] + ")")

            elif choice == 2:
                run_answ = paragraph_answ.add_run("Poprawna odp. : C (" + rest_answers[2] + ")")

            elif choice == 3:
                run_answ = paragraph_answ.add_run("Poprawna odp. : D (" + rest_answers[2] + ")")

            run_answ.font.name = 'Arial'
            run_answ.font.size = docx.shared.Pt(12)

        doc.save(docx_path)
        doc_answers.save(docx_path_answers)


def window():
    global cursor
    global connection

    app = QApplication(sys.argv)
    connection = sqlite3.connect("test-gen-db.db")
    create_new = False # WARNING: only change to True on first run of program!
    create_dir = False
    if check_db_connection(connection):
        cursor = connection.cursor()
        if create_new:
            create_new_tables()
        
        if create_dir:
            os.mkdir(path) 
        
        connection.close()

        win = MainAppWindow()
        win.show()
    else:
        print("Brak połączenia z bazą danych")

    
    sys.exit(app.exec())

window()